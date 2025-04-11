from docx import Document
from docx.shared import Inches, Pt
import geopandas as gpd
import rasterio
from rasterio.plot import show
import matplotlib.pyplot as plt
import numpy as np
from matplotlib_scalebar.scalebar import ScaleBar
import cartopy.crs as ccrs
from cartopy.mpl.ticker import LatitudeFormatter, LongitudeFormatter
from component.scripts.mosaics_helper import *
from datetime import datetime
import cartopy.io.img_tiles as cimgt
from shapely.geometry import box
from pathlib import Path
from component.parameter import directory


def ensure_list(input_data):
    if isinstance(input_data, str):
        return [input_data]  # Convert string to list
    elif isinstance(input_data, list):
        return input_data  # Return the list as is
    else:
        raise ValueError("Input must be a string or a list")  # Handle invalid input


def format_list(items):
    if not items:  # Handle empty list
        return ""
    if len(items) == 1:
        return str(items[0])
    elif len(items) == 2:
        return f"{items[0]} and {items[1]}"
    else:
        return f"{', '.join(map(str, items[:-1]))} and {items[-1]}"


def parse_formatted_string(formatted_string):
    if not formatted_string:  # Handle empty string
        return []
    if " and " not in formatted_string:  # Only one item
        return [formatted_string]
    parts = formatted_string.rsplit(" and ", 1)
    if ", " in parts[0]:  # Multiple items
        return parts[0].split(", ") + [parts[1]]
    else:  # Two items
        return parts


def plot_tiff_with_overlay(
    tiff_path, output_path, bands=(1, 2, 3), vector_overlay=None, overlay_color="red"
):
    """
    Plot a multi-band TIFF image with map elements and an optional vector overlay.

    Parameters:
    - tiff_path (str): Path to the input TIFF image.
    - output_path (str): Path to save the output JPEG image.
    - bands (tuple): Tuple of band indices for RGB visualization (e.g., (1, 2, 3) for Red, Green, Blue).
    - vector_overlay (GeoDataFrame): Optional GeoDataFrame for overlaying vector data on the map.
    - overlay_color (str): Color for vector overlay.
    """

    # Load the TIFF image
    with rasterio.open(tiff_path) as src:
        # Read the specified bands
        img_data = src.read(bands)
        img_extent = [
            src.bounds.left,
            src.bounds.right,
            src.bounds.bottom,
            src.bounds.top,
        ]
        crs_proj = src.crs.to_string()

    # Set up the plot
    fig, ax = plt.subplots(
        figsize=(6, 6), subplot_kw={"projection": ccrs.epsg(int(src.crs.to_epsg()))}
    )
    ax.set_extent(img_extent, crs=ccrs.epsg(int(src.crs.to_epsg())))

    # Display image data
    ax.imshow(img_data.transpose(1, 2, 0), extent=img_extent, origin="upper")

    # Add gridlines and ticks
    gl = ax.gridlines(draw_labels=True, color="gray", alpha=0.5, linestyle="--")
    gl.top_labels = True
    gl.bottom_labels = True
    gl.left_labels = True
    gl.right_labels = True

    # Rotate left and right tick labels for vertical orientation
    # gl.xlabel_style = {'rotation': 0}
    # gl.ylabel_style = {'rotation': 90}

    # Add north arrow
    ax.annotate(
        "N",
        xy=(0.95, 0.95),
        xycoords="axes fraction",
        ha="center",
        va="center",
        fontsize=16,
        color="black",
        weight="bold",
        arrowprops=dict(
            facecolor="black", width=4, headwidth=8, headlength=4, shrink=0.4
        ),
    )

    # Add scale bar
    # scale_length_km = 5  # scale length in km (customize as needed)
    # scale_length_deg = scale_length_km / 111  # approximate conversion to degrees
    # x0, y0 = ax.get_xlim()[0] + scale_length_deg * 1, ax.get_ylim()[0] + scale_length_deg * 1
    # ax.plot([x0, x0 + scale_length_deg], [y0, y0], color='black', lw=3)
    # ax.text(x0 + scale_length_deg / 2, y0, f'{scale_length_km} km', ha='center', va='bottom')
    ax.add_artist(
        ScaleBar(
            4.77,
            location="lower left",
            label_loc="bottom",
            scale_loc="top",
            frameon=True,
            length_fraction=0.1,
            width_fraction=0.004,
            box_alpha=0.2,
        )
    )

    # Add vector overlay if provided
    if vector_overlay is not None and not vector_overlay.empty:
        vector_overlay.to_crs(crs_proj, inplace=True)
        vector_overlay.plot(
            ax=ax,
            edgecolor=overlay_color,
            facecolor="none",
            linewidth=1,
            transform=ccrs.epsg(int(src.crs.to_epsg())),
        )

    # Save the output as JPEG
    plt.savefig(output_path, format="jpg", bbox_inches="tight", dpi=150)
    plt.close(fig)


# Check https://coolum001.github.io/cartopylayout.html
def add_north_arrow(ax, x=0.95, y=0.95):
    """Adds a north arrow to a Matplotlib axis."""
    ax.annotate(
        "N",
        xy=(x, y),
        xytext=(x, y - 0.05),
        xycoords="axes fraction",
        arrowprops=dict(
            facecolor="black", width=4, headwidth=8, headlength=4, shrink=0.4
        ),
        ha="center",
        fontsize=10,
    )


def generate_deforestation_report_with_word_template(
    image_path, geodataframe_path, template_path, output_path, scale_bar_length=200
):
    # Load Word template
    doc = Document(template_path)
    geodataframe = gpd.read_file(geodataframe_path)

    # Extract the first row's attributes from the GeoDataFrame
    gdf_row = geodataframe.iloc[0]  # Assuming one row of relevant data
    attributes = {
        "admin1": gdf_row["admin1"],
        "admin2": gdf_row["admin2"],
        "admin3": gdf_row["admin3"],
        "detection_date1": convert_julian_to_date(gdf_row["alert_date_min"]),
        "detection_date2": convert_julian_to_date(gdf_row["alert_date_max"]),
        "confirmation_date": datetime.now().strftime("%Y-%m-%d"),
        "before_img": gdf_row["before_img"],
        "after_img": gdf_row["after_img"],
        "alert_system": gdf_row["alert_sources"],
        "area_loss": "{:.2f}".format(gdf_row["area_ha"]),
    }

    # Replace placeholders in the Word document
    for para in doc.paragraphs:
        for key, value in attributes.items():
            if f"{{{key}}}" in para.text:
                para.text = para.text.replace(f"{{{key}}}", str(value))

    # Process Image 1
    folder_temp = os.path.abspath(directory.module_dir / "temp")
    img1_path_jpg = folder_temp + "/output_image1.jpg"
    plot_tiff_with_overlay(
        image_path,
        Path(img1_path_jpg),
        bands=(4, 3, 2),
        vector_overlay=None,
        overlay_color="red",
    )

    # Process Image 2
    img2_path_jpg = folder_temp + "/output_image2.jpg"
    plot_tiff_with_overlay(
        image_path,
        Path(img2_path_jpg),
        bands=(8, 7, 6),
        vector_overlay=None,
        overlay_color="red",
    )

    # Process Image 3
    img3_path_jpg = folder_temp + "/output_image3.jpg"
    plot_tiff_with_overlay(
        image_path,
        Path(img3_path_jpg),
        bands=(8, 7, 6),
        vector_overlay=geodataframe,
        overlay_color="red",
    )

    # Process Image 1
    folder_temp = os.path.abspath(directory.module_dir / "temp")
    img4_path_jpg = folder_temp + "/output_image4.jpg"
    plot_tiff_with_overlay(
        image_path,
        Path(img4_path_jpg),
        bands=(3, 2, 1),
        vector_overlay=None,
        overlay_color="red",
    )

    # Process Image 2
    img5_path_jpg = folder_temp + "/output_image5.jpg"
    plot_tiff_with_overlay(
        image_path,
        Path(img5_path_jpg),
        bands=(7, 6, 5),
        vector_overlay=None,
        overlay_color="red",
    )

    # Process Image 2
    img6_path_jpg = folder_temp + "/output_image6.jpg"
    plot_tiff_with_overlay(
        image_path,
        Path(img6_path_jpg),
        bands=(7, 6, 5),
        vector_overlay=geodataframe,
        overlay_color="red",
    )

    # Insert images in place of placeholders
    for para in doc.paragraphs:
        if "[Placeholder for Image 1]" in para.text:
            para.text = ""  # Clear placeholder text
            run = para.add_run()
            run.add_picture(img1_path_jpg, width=Pt(250))

        elif "[Placeholder for Image 2]" in para.text:
            para.text = ""  # Clear placeholder text
            run = para.add_run()
            run.add_picture(img2_path_jpg, width=Pt(250))

        elif "[Placeholder for Image 3]" in para.text:
            para.text = ""  # Clear placeholder text
            run = para.add_run()
            run.add_picture(img3_path_jpg, width=Pt(250))

        if "[Placeholder for Image 4]" in para.text:
            para.text = ""  # Clear placeholder text
            run = para.add_run()
            run.add_picture(img4_path_jpg, width=Pt(250))

        elif "[Placeholder for Image 5]" in para.text:
            para.text = ""  # Clear placeholder text
            run = para.add_run()
            run.add_picture(img5_path_jpg, width=Pt(250))

        elif "[Placeholder for Image 6]" in para.text:
            para.text = ""  # Clear placeholder text
            run = para.add_run()
            run.add_picture(img6_path_jpg, width=Pt(250))

    # Save the report as a Word document
    doc.save(output_path)
    # print(f"Report generated successfully: {output_path}")


def get_unique_alerts(alert_list):
    """
    Given a list of integers representing sums of alert values,
    return a list of unique alert types found.

    The alert values are constructed by summing contributions from:
    - GLAD-L: ones digit (1 or 2 if triggered)
    - RADD: tens digit (10 or 20 → digit 1 or 2 if triggered)
    - GLAD-S2: hundreds digit (100 or 200 → digit 1 or 2 if triggered)
    - CCDC: thousands digit (1000 or 2000 → digit 1 or 2 if triggered)

    Parameters:
        alert_list (list): A list of integers (or strings representing numbers).

    Returns:
        List[str]: Unique alert type names found in the list.
    """
    if not isinstance(alert_list, list):
        alert_list = [alert_list]
    
    unique_alerts = set()

    for value in alert_list:
        # Ensure value is an integer
        try:
            value = int(value)  # Convert to int if it's a string
        except ValueError:
            continue  # Skip invalid entries

        # Check GLAD-L (ones digit)
        if value % 10 in [1, 2]:
            unique_alerts.add("GLAD-L")
        
        # Check RADD (tens digit)
        if (value % 100) // 10 in [1, 2]:
            unique_alerts.add("RADD")
        
        # Check GLAD-S2 (hundreds digit)
        if (value % 1000) // 100 in [1, 2]:
            unique_alerts.add("GLAD-S2")
        
        # Check CCDC (thousands digit)
        if value // 1000 in [1, 2]:
            unique_alerts.add("CCDC")

    return list(unique_alerts)

